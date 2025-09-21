#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
DOCX Parser Module

This module provides functionality for parsing DOCX resume files.
"""

import os
from typing import Optional

from resume_analyzer.parsers.base_parser import BaseParser


class DocxParser(BaseParser):
    """
    Parser for DOCX resume files.
    Uses python-docx to extract text from DOCX files.
    """
    
    def __init__(self):
        """
        Initialize the DOCX parser.
        Tries to import required libraries and raises ImportError if they are not available.
        """
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx is not installed. "
                "Please install it using pip: "
                "pip install python-docx"
            )
    
    def parse(self, file_path: str) -> str:
        """
        Parse the DOCX file and extract its text content.
        
        Args:
            file_path (str): Path to the DOCX file.
            
        Returns:
            str: The extracted text content from the DOCX.
            
        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file cannot be parsed.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not os.path.isfile(file_path):
            raise ValueError(f"Not a file: {file_path}")
        
        try:
            return self._parse_with_docx(file_path)
        except Exception as e:
            raise ValueError(f"Error parsing DOCX file: {e}")
    
    def _parse_with_docx(self, file_path: str) -> str:
        """
        Parse the DOCX file using python-docx.
        
        Args:
            file_path (str): Path to the DOCX file.
            
        Returns:
            str: The extracted text content from the DOCX.
        """
        import docx
        
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text.append(para.text)
        
        return '\n'.join(full_text)